from __future__ import annotations

import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


@dataclass(frozen=True)
class ExtractedChunk:
    text: str
    page_number: int | None = None
    sheet_name: str = ""


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".xlsm",
    ".txt",
    ".md",
    ".png",
    ".jpg",
    ".jpeg",
}


def chunk_text(text: str, max_chars: int = 1800, overlap: int = 180) -> list[str]:
    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not normalized:
        return []
    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + max_chars, len(normalized))
        window = normalized[start:end]
        split_at = max(window.rfind("\n\n"), window.rfind("。"), window.rfind("\n"))
        if split_at > max_chars * 0.45 and end < len(normalized):
            end = start + split_at + 1
        chunks.append(normalized[start:end].strip())
        if end >= len(normalized):
            break
        start = max(0, end - overlap)
    return [part for part in chunks if part]


def _as_chunks(text: str, *, page_number: int | None = None, sheet_name: str = "") -> list[ExtractedChunk]:
    return [
        ExtractedChunk(text=part, page_number=page_number, sheet_name=sheet_name)
        for part in chunk_text(text)
    ]


class OcrEngine:
    def __init__(self, language: str = "ch") -> None:
        self.language = language
        self._engine = None

    def _load(self):
        if self._engine is None:
            from paddleocr import PaddleOCR

            try:
                self._engine = PaddleOCR(
                    lang=self.language,
                    ocr_version="PP-OCRv4",
                    use_textline_orientation=False,
                    use_doc_orientation_classify=False,
                    use_doc_unwarping=False,
                )
            except ValueError:
                self._engine = PaddleOCR(use_angle_cls=True, lang=self.language)
        return self._engine

    def image_to_text(self, image_path: Path) -> str:
        engine = self._load()
        if hasattr(engine, "predict"):
            result = engine.predict(str(image_path))
        else:
            result = engine.ocr(str(image_path), cls=True)
        return "\n".join(_collect_ocr_texts(result)).strip()


def _collect_ocr_texts(result: object) -> list[str]:
    texts: list[str] = []
    if result is None:
        return texts
    if isinstance(result, dict):
        rec_texts = result.get("rec_texts")
        if isinstance(rec_texts, list):
            texts.extend(str(text) for text in rec_texts if str(text).strip())
        return texts
    if isinstance(result, (list, tuple)):
        for item in result:
            if isinstance(item, dict):
                texts.extend(_collect_ocr_texts(item))
                continue
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                # PaddleOCR 2.x shape: [[box, (text, score)], ...]
                maybe_text = item[1]
                if isinstance(maybe_text, (list, tuple)) and maybe_text:
                    text = str(maybe_text[0]).strip()
                    if text:
                        texts.append(text)
                    continue
            texts.extend(_collect_ocr_texts(item))
    return texts


def extract_text_file(path: Path) -> list[ExtractedChunk]:
    return _as_chunks(path.read_text(encoding="utf-8-sig", errors="replace"))


def extract_docx(path: Path) -> list[ExtractedChunk]:
    from docx import Document

    document = Document(path)
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _as_chunks("\n".join(parts))


def extract_xlsx(path: Path) -> list[ExtractedChunk]:
    from openpyxl import load_workbook

    workbook = load_workbook(path, data_only=True, read_only=True)
    chunks: list[ExtractedChunk] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
        for part in chunk_text("\n".join(rows), max_chars=1600, overlap=120):
            chunks.append(ExtractedChunk(text=part, sheet_name=sheet.title))
    return chunks


def extract_image(path: Path, *, ocr_enabled: bool, ocr_language: str) -> tuple[list[ExtractedChunk], str]:
    if not ocr_enabled:
        return [ExtractedChunk(text="图片文件未启用 OCR，需人工复核。")], "required"
    text = OcrEngine(ocr_language).image_to_text(path)
    if not text:
        return [ExtractedChunk(text="图片 OCR 未识别到文本，需人工复核。")], "empty"
    return _as_chunks(text), "completed"


def _render_pdf_page_to_png(pdf_doc, page_index: int, target: Path) -> None:
    page = pdf_doc.load_page(page_index)
    matrix = __import__("fitz").Matrix(2, 2)
    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
    pixmap.save(target)


def extract_pdf(path: Path, *, ocr_enabled: bool, ocr_language: str) -> tuple[list[ExtractedChunk], str]:
    import pdfplumber

    chunks: list[ExtractedChunk] = []
    pages_needing_ocr: list[int] = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
            if text:
                chunks.extend(_as_chunks(text, page_number=index))
            else:
                pages_needing_ocr.append(index)

    if not pages_needing_ocr:
        return chunks, "not_needed"

    if not ocr_enabled:
        chunks.append(
            ExtractedChunk(
                text=f"PDF 第 {', '.join(map(str, pages_needing_ocr))} 页未抽取到文本，需人工 OCR 复核。"
            )
        )
        return chunks, "required"

    ocr = OcrEngine(ocr_language)
    ocr_text_found = False
    try:
        import fitz

        with fitz.open(path) as pdf_doc, tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            for page_no in pages_needing_ocr:
                image_path = temp_path / f"page_{page_no}.png"
                _render_pdf_page_to_png(pdf_doc, page_no - 1, image_path)
                text = ocr.image_to_text(image_path)
                if text:
                    ocr_text_found = True
                    chunks.extend(_as_chunks(text, page_number=page_no))
    except Exception as exc:
        chunks.append(ExtractedChunk(text=f"OCR 执行失败：{exc}。请人工复核扫描页。"))
        return chunks, "failed"

    return chunks, "completed" if ocr_text_found else "empty"


def chunks_to_dicts(chunks: Iterable[ExtractedChunk]) -> list[dict[str, object]]:
    return [
        {
            "text": chunk.text,
            "page_number": chunk.page_number,
            "sheet_name": chunk.sheet_name,
        }
        for chunk in chunks
        if chunk.text.strip()
    ]


def extract_document(
    path: Path,
    *,
    ocr_enabled: bool,
    ocr_language: str,
) -> tuple[list[dict[str, object]], str]:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix}")
    if suffix in {".txt", ".md"}:
        return chunks_to_dicts(extract_text_file(path)), "not_needed"
    if suffix == ".docx":
        return chunks_to_dicts(extract_docx(path)), "not_needed"
    if suffix in {".xlsx", ".xlsm"}:
        return chunks_to_dicts(extract_xlsx(path)), "not_needed"
    if suffix == ".pdf":
        chunks, ocr_status = extract_pdf(path, ocr_enabled=ocr_enabled, ocr_language=ocr_language)
        return chunks_to_dicts(chunks), ocr_status
    if suffix in {".png", ".jpg", ".jpeg"}:
        chunks, ocr_status = extract_image(path, ocr_enabled=ocr_enabled, ocr_language=ocr_language)
        return chunks_to_dicts(chunks), ocr_status
    raise ValueError(f"Unsupported file type: {suffix}")
